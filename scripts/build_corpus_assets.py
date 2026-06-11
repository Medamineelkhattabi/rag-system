"""Génère les documents binaires du corpus de démonstration (PDF et DOCX).

- data/vision_par_ordinateur.pdf    (via reportlab)
- data/guide_prompt_engineering.docx (via python-docx)

Contenu pédagogique, fictif et généraliste (uniquement sur l'IA). Relancer
après modification :
    python -m scripts.build_corpus_assets
"""
from __future__ import annotations

from pathlib import Path

DATA_DIR = Path(__file__).resolve().parent.parent / "data"


# PDF : introduction à la vision par ordinateur
PDF_BLOCKS: list[tuple[str, str]] = [
    ("h1", "Introduction à la vision par ordinateur"),
    ("i", "Document pédagogique de démonstration, fictif et généraliste."),
    ("h2", "1. Définition"),
    ("p", "La vision par ordinateur est le domaine de l'intelligence artificielle "
          "qui permet à une machine d'analyser et d'interpréter des images ou des "
          "vidéos. L'objectif est d'extraire automatiquement de l'information "
          "visuelle : identifier des objets, lire du texte, mesurer, ou décrire "
          "une scène."),
    ("h2", "2. Traitement d'image"),
    ("p", "Le traitement d'image regroupe les opérations de base appliquées aux "
          "images : redimensionnement, normalisation, ajustement du contraste, "
          "réduction du bruit et détection de contours. Ces étapes de "
          "prétraitement améliorent la qualité des données avant l'analyse par un "
          "modèle."),
    ("h2", "3. Classification d'image"),
    ("p", "La classification consiste à attribuer une étiquette globale à une "
          "image, par exemple déterminer si une photo représente un chat ou un "
          "chien. Le modèle apprend à associer des motifs visuels à des "
          "catégories à partir d'exemples étiquetés."),
    ("h2", "4. Détection d'objets"),
    ("p", "La détection d'objets localise et identifie plusieurs éléments dans une "
          "même image, en traçant des boîtes englobantes autour de chacun. Elle "
          "répond à la question : quels objets sont présents et où se trouvent-ils ?"),
    ("h2", "5. Segmentation"),
    ("p", "La segmentation va plus loin que la détection : elle attribue une "
          "catégorie à chaque pixel de l'image. On distingue la segmentation "
          "sémantique (par classe) et la segmentation d'instances (objet par "
          "objet). Elle est utile lorsque la forme précise des objets compte."),
    ("h2", "6. Reconnaissance optique de caractères (OCR)"),
    ("p", "L'OCR convertit le texte présent dans une image (document scanné, "
          "panneau, étiquette) en texte exploitable par une machine. Elle combine "
          "souvent détection des zones de texte et reconnaissance des caractères."),
    ("h2", "7. Reconnaissance faciale (exemple théorique)"),
    ("p", "La reconnaissance faciale est présentée ici uniquement comme exemple "
          "théorique, sans aucune donnée personnelle ni donnée biométrique réelle. "
          "Sur le principe, le système détecte un visage dans une image puis "
          "compare une représentation numérique à une référence. Ce type d'usage "
          "soulève d'importantes questions de vie privée et de réglementation, qui "
          "doivent être traitées avec rigueur."),
    ("h2", "8. Exemples industriels"),
    ("p", "Le contrôle qualité visuel sur une chaîne de production, le comptage "
          "d'objets, l'inspection d'infrastructures à partir d'images de drones, "
          "ou encore l'assistance au tri sont des exemples d'applications "
          "industrielles de la vision par ordinateur."),
    ("h2", "9. Limites"),
    ("p", "Les performances dépendent fortement des conditions : éclairage, angle "
          "de prise de vue, qualité et résolution de l'image. Les modèles peuvent "
          "aussi reproduire des biais présents dans les données et manquer de "
          "robustesse face à des situations nouvelles ou dégradées."),
    ("h2", "10. Lien avec les modèles vision-langage (VLM)"),
    ("p", "Un modèle vision-langage (VLM) combine la compréhension des images et "
          "du texte. Il peut décrire une image en langage naturel, répondre à des "
          "questions sur son contenu ou relier une consigne textuelle à des "
          "éléments visuels. Les VLM étendent ainsi la vision par ordinateur vers "
          "des usages multimodaux, à la frontière avec l'IA générative."),
]


def build_pdf(path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    flow = []
    for kind, text in PDF_BLOCKS:
        if kind == "h1":
            flow.append(Paragraph(text, styles["Title"]))
        elif kind == "h2":
            flow.append(Paragraph(text, styles["Heading2"]))
        elif kind == "i":
            flow.append(Paragraph(f"<i>{text}</i>", styles["Normal"]))
        else:
            flow.append(Paragraph(text, styles["BodyText"]))
        flow.append(Spacer(1, 8))
    SimpleDocTemplate(str(path), pagesize=A4, title="Vision par ordinateur").build(flow)


# DOCX : guide de prompt engineering
DOCX_BLOCKS: list[tuple[str, str]] = [
    ("h1", "Guide de prompt engineering"),
    ("i", "Document pédagogique de démonstration, fictif et généraliste."),
    ("h2", "1. Qu'est-ce que le prompt engineering ?"),
    ("p", "Le prompt engineering est l'art de formuler des consignes (prompts) "
          "efficaces pour obtenir d'un modèle de langage des réponses utiles, "
          "précises et fiables. Un bon prompt guide le modèle sur la tâche, le "
          "format attendu et le niveau de détail souhaité."),
    ("h2", "2. Importance du contexte"),
    ("p", "Plus le contexte fourni est pertinent, meilleure est la réponse. "
          "Donner des informations de fond, des exemples ou des documents de "
          "référence aide le modèle à rester factuel et à coller au besoin réel."),
    ("h2", "3. Formuler une consigne claire"),
    ("p", "Une bonne consigne est explicite : elle précise le rôle attendu, la "
          "tâche, les contraintes et le format de sortie. Il vaut mieux décomposer "
          "une demande complexe en étapes claires que de tout regrouper en une "
          "phrase ambiguë."),
    ("h2", "4. Mauvais prompt et bon prompt"),
    ("p", "Mauvais prompt : « Parle-moi de ce texte. » Bon prompt : « Résume ce "
          "texte en cinq points clés, en français, destinés à un lecteur non "
          "spécialiste. » Le second précise la tâche, le format, la langue et le "
          "public, ce qui réduit l'ambiguïté."),
    ("h2", "5. Prompt pour résumer un document"),
    ("p", "Exemple : « Résume le document suivant en un paragraphe de cinq lignes "
          "maximum, en conservant uniquement les idées principales. Document : "
          "[texte]. » On précise la longueur et l'objectif."),
    ("h2", "6. Prompt pour extraire des informations"),
    ("p", "Exemple : « À partir du texte ci-dessous, extrais les dates, les "
          "montants et les noms d'organisations, et présente-les sous forme de "
          "liste structurée. » On indique précisément ce qu'il faut extraire et "
          "le format attendu."),
    ("h2", "7. Prompt pour comparer deux concepts"),
    ("p", "Exemple : « Compare l'apprentissage supervisé et l'apprentissage non "
          "supervisé sous forme de tableau, avec les colonnes : définition, "
          "données utilisées, exemple d'application. » La structure demandée rend "
          "la réponse plus exploitable."),
    ("h2", "8. Prompt pour générer du code"),
    ("p", "Exemple : « Écris une fonction Python qui calcule la similarité cosinus "
          "entre deux vecteurs, avec un commentaire expliquant chaque étape et un "
          "exemple d'utilisation. » On précise le langage, la tâche et les "
          "attentes de documentation."),
    ("h2", "9. Prompt pour demander des sources"),
    ("p", "Exemple : « Réponds uniquement à partir du contexte fourni et indique, "
          "pour chaque affirmation, la source correspondante. Si l'information est "
          "absente du contexte, dis-le explicitement. » C'est une consigne clé "
          "pour un système RAG fiable."),
    ("h2", "10. Bonnes pratiques pour des réponses précises"),
    ("p", "Être spécifique, fournir des exemples, préciser le format de sortie, "
          "limiter la longueur attendue, demander un raisonnement étape par étape "
          "pour les tâches complexes, et itérer en ajustant le prompt selon les "
          "résultats obtenus."),
    ("h2", "11. Limites du prompt engineering"),
    ("p", "Le prompt engineering ne corrige pas tout : il ne remplace pas des "
          "données fiables ni une architecture adaptée (par exemple un RAG pour "
          "ancrer les réponses). Un modèle peut toujours se tromper malgré un bon "
          "prompt ; la vérification des informations importantes reste nécessaire."),
]


def build_docx(path: Path) -> None:
    import docx

    document = docx.Document()
    for kind, text in DOCX_BLOCKS:
        if kind == "h1":
            document.add_heading(text, level=0)
        elif kind == "h2":
            document.add_heading(text, level=1)
        elif kind == "i":
            p = document.add_paragraph()
            run = p.add_run(text)
            run.italic = True
        else:
            document.add_paragraph(text)
    document.save(str(path))


def main() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    pdf_path = DATA_DIR / "vision_par_ordinateur.pdf"
    docx_path = DATA_DIR / "guide_prompt_engineering.docx"
    build_pdf(pdf_path)
    build_docx(docx_path)
    print(f"[OK] PDF genere  : {pdf_path}")
    print(f"[OK] DOCX genere : {docx_path}")


if __name__ == "__main__":
    main()
